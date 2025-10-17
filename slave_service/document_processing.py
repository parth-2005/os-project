"""Document processing module for analyzing document structure and metadata"""
import base64
import json
from PyPDF2 import PdfReader
import docx

def process_documents(doc_files):
    """
    Process document files (PDF, DOCX) to extract metadata, text, and structure.
    Returns document analysis for each file.
    """
    results = []
    for doc_file in doc_files:
        try:
            filename = doc_file.filename.lower()
            
            if filename.endswith('.pdf'):
                doc_analysis = process_pdf(doc_file)
            elif filename.endswith('.docx'):
                doc_analysis = process_docx(doc_file)
            else:
                # Generic text extraction for other formats
                doc_analysis = process_generic_doc(doc_file)
            
            # Encode as base64 JSON
            result_str = base64.b64encode(json.dumps(doc_analysis).encode()).decode('utf-8')
            
            results.append({
                "filename": doc_file.filename,
                "document_data": result_str
            })
        except Exception as e:
            print(f"Error processing document {doc_file.filename}: {e}")
    
    return results

def process_pdf(pdf_file):
    """Extract text and metadata from PDF"""
    try:
        reader = PdfReader(pdf_file.stream)
        
        # Extract text from all pages
        text_content = ""
        for page in reader.pages:
            text_content += page.extract_text()
        
        # Get metadata
        metadata = reader.metadata or {}
        
        return {
            "document_type": "PDF",
            "page_count": len(reader.pages),
            "text_content": text_content[:1000],  # First 1000 chars
            "word_count": len(text_content.split()),
            "character_count": len(text_content),
            "metadata": {
                "title": str(metadata.get('/Title', '')),
                "author": str(metadata.get('/Author', '')),
                "subject": str(metadata.get('/Subject', '')),
                "creator": str(metadata.get('/Creator', ''))
            }
        }
    except Exception as e:
        return {"document_type": "PDF", "error": str(e)}

def process_docx(docx_file):
    """Extract text and metadata from DOCX"""
    try:
        doc = docx.Document(docx_file.stream)
        
        # Extract text from paragraphs
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Count elements
        table_count = len(doc.tables)
        
        return {
            "document_type": "DOCX",
            "paragraph_count": len(doc.paragraphs),
            "table_count": table_count,
            "text_content": text_content[:1000],  # First 1000 chars
            "word_count": len(text_content.split()),
            "character_count": len(text_content)
        }
    except Exception as e:
        return {"document_type": "DOCX", "error": str(e)}

def process_generic_doc(doc_file):
    """Generic text extraction for other document types"""
    try:
        content = doc_file.stream.read()
        
        # Try to decode as text
        try:
            text_content = content.decode('utf-8')
        except:
            text_content = content.decode('utf-8', errors='ignore')
        
        return {
            "document_type": "Generic Text",
            "text_content": text_content[:1000],  # First 1000 chars
            "word_count": len(text_content.split()),
            "character_count": len(text_content),
            "file_size_bytes": len(content)
        }
    except Exception as e:
        return {"document_type": "Generic", "error": str(e)}